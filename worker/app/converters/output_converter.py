import os
import re

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt
from PIL import Image


def save_markdown(page_texts: list[str], output_path: str) -> str:
    """Save OCR results as a Markdown file."""
    content = "\n\n---\n\n".join(page_texts)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)
    return output_path


def save_plaintext(page_texts: list[str], output_path: str) -> str:
    """Save OCR results as plain text, stripping markdown formatting."""
    cleaned_pages = []
    for text in page_texts:
        clean = _strip_markdown(text)
        cleaned_pages.append(clean)
    content = "\n\n" + ("=" * 40) + "\n\n".join(cleaned_pages)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)
    return output_path


def save_searchable_pdf(
    original_path: str,
    page_texts: list[str],
    output_path: str,
) -> str:
    """Create a searchable PDF with invisible text overlay.

    For PDF inputs: overlays text on original pages.
    For image inputs: creates a new PDF from the image(s) with text overlay.
    """
    ext = os.path.splitext(original_path)[1].lower()

    if ext == ".pdf":
        doc = fitz.open(original_path)
    else:
        # Create a PDF from the image
        doc = fitz.open()
        img = Image.open(original_path)
        # Handle multi-frame images
        frames = [img]
        try:
            while True:
                img.seek(img.tell() + 1)
                frames.append(img.copy())
        except EOFError:
            pass

        for frame in frames:
            if frame.mode != "RGB":
                frame = frame.convert("RGB")
            import io
            buf = io.BytesIO()
            frame.save(buf, format="PNG")
            buf.seek(0)
            img_rect = fitz.Rect(0, 0, frame.width, frame.height)
            page = doc.new_page(width=frame.width, height=frame.height)
            page.insert_image(img_rect, stream=buf.getvalue())

    # Overlay invisible text on each page
    for page_num, text in enumerate(page_texts):
        if page_num >= len(doc):
            break
        page = doc[page_num]
        # Use a very small font and render_mode=3 (invisible)
        # Insert as a text block covering the page
        rc = page.insert_textbox(
            page.rect,
            text,
            fontsize=8,
            fontname="helv",
            render_mode=3,  # invisible text
        )
        # rc < 0 means overflow, which is fine for searchable text

    doc.save(output_path)
    doc.close()
    return output_path


def save_docx(page_texts: list[str], output_path: str) -> str:
    """Save OCR results as a Word document using Pandoc for proper table conversion."""
    import pypandoc

    content = "\n\n---\n\n".join(page_texts)
    pypandoc.convert_text(
        content,
        "docx",
        format="markdown",
        outputfile=output_path,
    )
    return output_path


def _strip_markdown(text: str) -> str:
    """Remove basic markdown formatting."""
    # Remove headers
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    # Remove bold/italic
    text = re.sub(r"\*{1,3}(.+?)\*{1,3}", r"\1", text)
    text = re.sub(r"_{1,3}(.+?)_{1,3}", r"\1", text)
    # Remove links
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    # Remove code ticks
    text = re.sub(r"`(.+?)`", r"\1", text)
    # Remove horizontal rules
    text = re.sub(r"^-{3,}$", "", text, flags=re.MULTILINE)
    return text.strip()


def _add_markdown_to_docx(doc: Document, text: str) -> None:
    """Convert basic markdown text into Word document elements."""
    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue

        # Detect headings
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            doc.add_heading(heading_match.group(2), level=level)
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        run = p.add_run(stripped)
        run.font.size = Pt(11)
